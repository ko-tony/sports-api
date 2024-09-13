from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
from openpyxl.drawing.image import Image as ExcelImage
from PIL import Image as PILImage
from io import BytesIO
import zipfile
from fastapi.responses import StreamingResponse
from typing import List, Tuple
from datetime import datetime

def excel_to_word(file_contents: bytes) -> StreamingResponse:
    
    # 读取上传的 xlsx 文件
    workbook = load_workbook(BytesIO(file_contents))
    sheet = workbook.active
    
    # 指定要抓取的关键词
    keywords = ["上線日", "文章類型", "預計版位", "文章標題", "開文內容", "開文圖一", "開文圖二", "開文圖三", "開文圖四", "回文1", "回文2", "回文3", "回文4",	"回文5", "回文6", "回文7", "回文8", "回文9", "回文10"]
    
    # 获取需要提取的行
    rows_to_extract = get_rows_to_extract(sheet, keywords)

    # 创建 ZIP 文件
    zip_io = create_zip_file(sheet, rows_to_extract)

    # 返回 StreamingResponse
    current_time = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"{current_time}_report.zip"

    return StreamingResponse(
        zip_io, 
        media_type="application/x-zip-compressed",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

def get_rows_to_extract(sheet, keywords: List[str]) -> List[int]:
    rows_to_extract = []
    for idx, row in enumerate(sheet.iter_rows(min_row=1, max_row=50, values_only=True), start=1):
        if any(keyword in str(cell) for cell in row for keyword in keywords if cell):
            rows_to_extract.append(idx)
    return rows_to_extract

def create_word_doc(sheet, rows_to_extract: List[int], col1: int, col2: int) -> bytes:
    document = Document()
    table = document.add_table(rows=1, cols=2)

    for row_idx in rows_to_extract:
        row = sheet[row_idx]
        row_cells = table.add_row().cells
        
        for col_idx, cell in enumerate(row_cells):
            excel_cell = row[col1] if col_idx == 0 else row[col2]
            if excel_cell.value is not None:
                cell.text = str(excel_cell.value)
            else:
                add_image_to_cell(sheet, cell, row_idx, col1 if col_idx == 0 else col2)

    docx_io = BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io.getvalue()

def add_image_to_cell(sheet, cell, row_idx: int, col_idx: int):
    for image in sheet._images:
        # row_idx-1 才不會漏行
        if image.anchor._from.row == row_idx-1 and image.anchor._from.col == col_idx:
            image_stream = BytesIO()
            pil_image = PILImage.open(BytesIO(image._data()))
            pil_image.save(image_stream, format='PNG' if pil_image.format == 'PNG' else 'JPEG')
            image_stream.seek(0)
            run = cell.paragraphs[0].add_run()
            run.add_picture(image_stream, width=Inches(1.25))
            break

def create_zip_file(sheet, rows_to_extract: List[int]) -> BytesIO:
    zip_io = BytesIO()
    with zipfile.ZipFile(zip_io, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for i in range(2, 21):
            if sheet.cell(row=1, column=i).value is None:
                break
            output_filename = f"report_{i-1}.docx"
            content = create_word_doc(sheet, rows_to_extract, 0, i-1)
            zip_file.writestr(output_filename, content)
    zip_io.seek(0)
    return zip_io